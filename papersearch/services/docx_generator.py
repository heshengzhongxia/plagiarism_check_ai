"""
DOCX report generator for plagiarism check reports
"""
import os
import time

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _set_font(run, font_name='宋体', size=None, bold=None, color=None):
    """统一设置中英文字体"""
    run.font.name = font_name
    # 设置中文字体（eastAsia），python-docx 不直接支持，需操作 XML
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def generate_docx(user_full_text, matches, modifications, overall_suggestions, report_dir, aigc_rate=0, aigc_analysis=""):
    """生成查重报告.docx文件"""
    if not HAS_DOCX:
        return None, "python-docx未安装"

    FONT_NAME = '宋体'

    doc = Document()

    # 全局默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_NAME)

    for i in range(3):
        hs = doc.styles[f'Heading {i+1}']
        hs.font.name = FONT_NAME
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_NAME)

    # 标题
    title = doc.add_heading('论文查重报告', 0)
    for run in title.runs:
        _set_font(run, FONT_NAME)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _run(text, bold=False, size=Pt(11), color=None):
        """创建格式统一的段落"""
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_font(r, FONT_NAME, size=size, bold=bold, color=color)
        return p

    _run(f"生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')} | 重复句子数：{len(matches)} | AIGC概率：{aigc_rate}%", size=Pt(10))
    if aigc_analysis:
        _run(f"AIGC分析：{aigc_analysis}", size=Pt(10))
    doc.add_paragraph("")

    # === 第一部分：用户论文全文 ===
    h = doc.add_heading('一、论文原文', 1)
    for r in h.runs: _set_font(r, FONT_NAME)
    for para_text in user_full_text.split('\n'):
        if para_text.strip():
            _run(para_text.strip())

    doc.add_page_break()

    # === 第二部分：逐句查重结果 ===
    h2 = doc.add_heading('二、逐句查重结果', 1)
    for r in h2.runs: _set_font(r, FONT_NAME)

    if not matches:
        _run("未检测到重复句子。", bold=True)
    else:
        for i, m in enumerate(matches):
            user_sent = m.get("user_sentence", "")
            similar = m.get("similar_sentence", "")
            source = m.get("source_title", "")
            url = m.get("source_url", "")
            sim = m.get("similarity", 0)

            h3 = doc.add_heading(f"重复句 {i+1}（相似度 {sim}%）", 2)
            for r in h3.runs: _set_font(r, FONT_NAME)

            _run(f"【用户原文】{user_sent}")

            p2 = _run(f"【相似句子】{similar}", color=RGBColor(200, 50, 50))
            p2.paragraph_format.left_indent = Cm(1.2)

            p3 = _run(f"来源：{source}" + (f"\n链接：{url}" if url else ""), size=Pt(10))
            p3.paragraph_format.left_indent = Cm(1.2)

            mod_text = ""
            for mod in modifications:
                if mod.get("user_sentence", "")[:40] == user_sent[:40]:
                    mod_text = mod.get("modified_sentence", "")
                    break
            if mod_text:
                p4 = _run(f"【修改建议】{mod_text}", color=RGBColor(0, 100, 0))
                p4.paragraph_format.left_indent = Cm(1.2)

            doc.add_paragraph("")

    doc.add_page_break()

    # === 第三部分：总体建议 ===
    h4 = doc.add_heading('三、总体修改建议', 1)
    for r in h4.runs: _set_font(r, FONT_NAME)

    if overall_suggestions:
        for s in overall_suggestions:
            priority = s.get("priority", "")
            title_s = s.get("title", "")
            content = s.get("content", "")
            _run(f"[{priority}] {title_s}", bold=True)
            _run(content)
    else:
        _run("暂无总体建议。")

    # 保存
    filename = f"查重报告_{int(time.time())}.docx"
    filepath = os.path.join(report_dir, filename)
    doc.save(filepath)
    return filepath, None
